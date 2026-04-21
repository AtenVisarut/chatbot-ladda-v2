"""
Generate _CROP_STAGES_DB.docx — reference document for PD team to review
crop stage data used by the bot's clarification-question system.

Contains:
  1. Current _CROP_STAGES mapping (from response_generator_agent.py)
  2. Thailand standard references for comparison
  3. Proposed days-to-stage mapping (not yet implemented)
  4. Action items + review checklist

Usage: python scripts/generate_crop_stages_doc.py
Output: reports/_CROP_STAGES_DB.docx
"""

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PRIMARY = RGBColor(0x0F, 0x4C, 0x81)
ACCENT = RGBColor(0x1F, 0x6F, 0xEB)
SUBTLE = RGBColor(0x55, 0x55, 0x55)
WARN = RGBColor(0xD9, 0x30, 0x25)


def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_banner(doc, title, subtitle=None, color="0F4C81"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if subtitle:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)
    doc.add_paragraph()


def add_paragraph(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.runs[0] if p.runs else p.add_run()
    r.text = text
    r.font.size = Pt(11)
    return p


# =============================================================================
# Data — current _CROP_STAGES from code
# =============================================================================

CURRENT_CROP_STAGES = [
    # category, crop, stages_in_code
    ("ไม้ผล (Fruit tree)", "ทุเรียน", "ใบอ่อน / แตกใบ / ออกดอก / ติดผล / ผลอ่อน / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ไม้ผล", "มะม่วง", "ใบอ่อน / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ไม้ผล", "ลำไย", "ใบอ่อน / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ไม้ผล", "ลิ้นจี่", "ใบอ่อน / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ไม้ผล", "ส้ม", "ใบอ่อน / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ไม้ผล", "ส้มโอ", "ใบอ่อน / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ไม้ผล", "มะนาว", "ใบอ่อน / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ไม้ผล", "เงาะ", "ใบอ่อน / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ไม้ผล", "มังคุด", "ใบอ่อน / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ปาล์ม", "ปาล์ม", "ต้นเล็ก / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("ปาล์ม", "ปาล์มน้ำมัน", "ต้นเล็ก / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("มะพร้าว", "มะพร้าว", "ต้นเล็ก / ออกดอก / ติดผล / ผลแก่ / หลังเก็บเกี่ยว"),
    ("พืชไร่ (Field crop)", "ข้าว", "ต้นกล้า / แตกกอ / ตั้งท้อง / ออกรวง / สุก / เก็บเกี่ยว"),
    ("พืชไร่", "ข้าวโพด", "ต้นกล้า / ก่อนออกดอก / ออกดอก / ติดฝัก / ฝักแก่"),
    ("พืชไร่", "อ้อย", "ปลูกใหม่ / แตกกอ / ยืดปล้อง / ก่อนเก็บเกี่ยว"),
    ("พืชไร่", "มันสำปะหลัง", "ปลูกใหม่ / ต้นเล็ก / สร้างหัว / ก่อนเก็บเกี่ยว"),
    ("ไม้ยืนต้น", "ยางพารา", "ต้นเล็ก / ให้น้ำยาง"),
    ("ผักใบ (Leaf veg)", "ผัก", "เพาะกล้า / ต้นอ่อน / โต / เก็บเกี่ยว"),
    ("ผักใบ", "ผักคะน้า", "เพาะกล้า / ต้นอ่อน / โต / เก็บเกี่ยว"),
    ("ผักใบ", "ผักกาด", "เพาะกล้า / ต้นอ่อน / โต / เก็บเกี่ยว"),
    ("ผักใบ", "ผักกาดขาว", "เพาะกล้า / ต้นอ่อน / โต / เก็บเกี่ยว"),
    ("ผักใบ", "คะน้า", "เพาะกล้า / ต้นอ่อน / โต / เก็บเกี่ยว"),
    ("ผักหัว", "กะหล่ำปลี", "เพาะกล้า / ต้นอ่อน / โต / เก็บเกี่ยว"),
    ("ผักผล", "พริก", "ต้นกล้า / ออกดอก / ติดผล / ผลแก่"),
    ("ผักผล", "มะเขือ", "ต้นกล้า / ออกดอก / ติดผล / ผลแก่"),
    ("ผักผล", "มะเขือเทศ", "ต้นกล้า / ออกดอก / ติดผล / ผลแก่"),
    ("ผักผล", "แตงกวา", "ต้นกล้า / ออกดอก / ติดผล / ผลแก่"),
    ("ถั่ว", "ถั่ว", "ต้นกล้า / ออกดอก / ติดฝัก / ฝักแก่"),
    ("ถั่ว", "ถั่วฝักยาว", "ต้นกล้า / ออกดอก / ติดฝัก / ฝักแก่"),
    ("ผักหัว (Bulb/root)", "หอม", "ต้นกล้า / แตกใบ / ลงหัว / เก็บเกี่ยว"),
    ("ผักหัว", "หอมแดง", "ต้นกล้า / แตกใบ / ลงหัว / เก็บเกี่ยว"),
    ("ผักหัว", "กระเทียม", "ต้นกล้า / แตกใบ / ลงหัว / เก็บเกี่ยว"),
    ("ผักหัว", "มันฝรั่ง", "ต้นกล้า / แตกใบ / ลงหัว / เก็บเกี่ยว"),
]


# Thailand standards reference for comparison
STANDARD_REFERENCES = [
    {
        "crop": "ข้าว",
        "source": "กรมการข้าว / Department of Rice",
        "standard": "5 ระยะหลัก",
        "stages": "Seedling (ต้นกล้า) → Tillering (แตกกอ) → Booting (ตั้งท้อง) → Heading (ออกรวง) → Maturity (สุก)",
        "match": "ใกล้เคียง ✓",
    },
    {
        "crop": "ข้าวโพด",
        "source": "มาตรฐานวิชาการ V/R scale (FAO)",
        "standard": "V1-V6 (Vegetative), VT (Tasseling), R1-R6 (Reproductive)",
        "stages": (
            "V1-V6 → VT (ออกดอกตัวผู้) → R1 (Silking - ติดฝัก) → "
            "R2 (Blister) → R3 (Milk) → R4 (Dough) → R5 (Dent) → R6 (Maturity)"
        ),
        "match": "ของเราใช้ง่ายกว่า — ควรเทียบกับมาตรฐาน V/R",
    },
    {
        "crop": "ทุเรียน",
        "source": "สมาคมทุเรียนไทย / กรมส่งเสริมการเกษตร",
        "standard": "ใช้คำสั่งงานจริง",
        "stages": (
            "แตกใบอ่อน (flushing) → ใบเพสลาด → ก่อนเปิดตา → ดอกบาน → "
            "ติดลูกเล็ก → ลูกขนุน → ผลแก่ (4 เดือน) → เก็บเกี่ยว"
        ),
        "match": "ของเราง่ายเกิน — ขาด ใบเพสลาด, ก่อนเปิดตา, ดอกบาน, ลูกขนุน",
    },
    {
        "crop": "มะม่วง",
        "source": "กรมวิชาการเกษตร",
        "standard": "ระยะตามรอบดอก",
        "stages": "แตกใบ → ก่อนดอก (ราดสาร) → ดอกบาน → ติดผลเล็ก → ลูกไข่ → ผลแก่",
        "match": "ใกล้เคียง แต่ขาด 'ราดสาร' (PGR application stage)",
    },
    {
        "crop": "อ้อย",
        "source": "กรมวิชาการเกษตร",
        "standard": "4 ระยะ",
        "stages": "ปลูก/แตกกอ (0-90 วัน) → ยืดปล้อง (90-180 วัน) → สะสมน้ำตาล (180-300 วัน) → เก็บเกี่ยว",
        "match": "ของเราขาด 'สะสมน้ำตาล'",
    },
    {
        "crop": "มันสำปะหลัง",
        "source": "กรมวิชาการเกษตร",
        "standard": "4 ระยะ",
        "stages": "ปลูก → แตกหน่อ (0-60d) → สะสมแป้ง (60-180d) → หัวแก่ (180-300d)",
        "match": "ของเราขาด 'สะสมแป้ง'",
    },
]


# Proposed days-to-stage mapping
DAYS_TO_STAGE = [
    {
        "crop": "ข้าว",
        "mapping": [
            ("0-20 วัน", "ต้นกล้า"),
            ("20-50 วัน", "แตกกอ"),
            ("50-75 วัน", "ตั้งท้อง"),
            ("75-90 วัน", "ออกรวง"),
            ("90-120 วัน", "สุก / เก็บเกี่ยว"),
        ],
    },
    {
        "crop": "ข้าวโพด",
        "mapping": [
            ("0-14 วัน", "ต้นกล้า"),
            ("14-40 วัน", "Vegetative (V1-V6)"),
            ("40-55 วัน", "ออกดอก (VT)"),
            ("55-70 วัน", "ติดฝัก (R1)"),
            ("70-110 วัน", "ฝักแก่ (R2-R6)"),
        ],
    },
    {
        "crop": "อ้อย",
        "mapping": [
            ("0-90 วัน", "แตกกอ"),
            ("90-180 วัน", "ยืดปล้อง"),
            ("180-300 วัน", "สะสมน้ำตาล"),
            (">300 วัน", "เก็บเกี่ยว"),
        ],
    },
    {
        "crop": "มันสำปะหลัง",
        "mapping": [
            ("0-60 วัน", "แตกหน่อ"),
            ("60-180 วัน", "สะสมแป้ง"),
            ("180-300 วัน", "หัวแก่ / เก็บเกี่ยว"),
        ],
    },
    {
        "crop": "ทุเรียน (รายเดือน)",
        "mapping": [
            ("หลังเปิดตา 7-30 วัน", "ดอกบาน"),
            ("หลังดอกบาน 30-60 วัน", "ติดลูกเล็ก"),
            ("หลังดอกบาน 60-100 วัน", "ลูกขนุน"),
            ("หลังดอกบาน 100-120 วัน", "ผลแก่"),
        ],
    },
    {
        "crop": "ผักใบ (คะน้า/กะหล่ำ)",
        "mapping": [
            ("0-15 วัน", "เพาะกล้า"),
            ("15-30 วัน", "ต้นอ่อน"),
            ("30-45 วัน", "โต"),
            (">45 วัน", "เก็บเกี่ยว"),
        ],
    },
]


# =============================================================================
# Document generation
# =============================================================================


def build(out_path: Path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("_CROP_STAGES_DB")
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ข้อมูลระยะพืช — สำหรับทีม PD review")
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"สร้างเมื่อ {datetime.now().strftime('%d %B %Y')}")
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = SUBTLE

    doc.add_paragraph()

    # --- Warning ---
    add_banner(
        doc,
        "⚠️ ข้อสำคัญ (Important)",
        "ข้อมูลชุดนี้เป็น best-effort approximation — ต้อง review โดยทีม PD ก่อนใช้งานจริง",
        color="D93025",
    )
    add_paragraph(
        doc,
        "• ข้อมูลระยะพืชทั้งหมดในเอกสารนี้มาจาก general knowledge ไม่ได้อ้างอิง "
        "จากเอกสารมาตรฐานของกรมวิชาการเกษตร/กรมการข้าว/สมาคมเฉพาะทาง",
    )
    add_paragraph(
        doc,
        "• ก่อนใช้งานใน production ต้องให้ทีม PD/เจ้าหน้าที่เกษตร ICP Ladda "
        "ยืนยันคำศัพท์ว่าตรงกับที่ใช้ภาคสนามจริง",
    )
    add_paragraph(
        doc,
        "• ระยะบางเรื่อง (ราดสาร, สะสมน้ำตาล, สะสมแป้ง, ใบเพสลาด, ลูกขนุน) "
        "ยังไม่ได้ใส่ใน system — ต้องขยายเมื่อมี requirement",
    )

    # --- Table 1: Current data in code ---
    doc.add_page_break()
    add_banner(
        doc,
        "1. ข้อมูลปัจจุบันใน Code",
        "มาจาก response_generator_agent.py — _CROP_STAGES dict",
    )
    add_paragraph(
        doc,
        "ตารางนี้คือข้อมูลที่ bot ใช้อยู่ตอนนี้ เพื่อสร้างประโยคถามระยะพืชเช่น "
        '"ระยะของทุเรียนตอนนี้ (เช่น ใบอ่อน/แตกใบ/...)"',
        italic=True, color=SUBTLE, size=10,
    )

    t = doc.add_table(rows=len(CURRENT_CROP_STAGES) + 1, cols=3)
    t.style = "Light Grid Accent 1"
    headers = ["ประเภท", "พืช", "ระยะใน code"]
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        set_cell_shading(c, "1F6FEB")
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (cat, crop, stages) in enumerate(CURRENT_CROP_STAGES, 1):
        row = t.rows[i].cells
        row[0].text = cat
        row[1].text = crop
        row[2].text = stages
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].paragraphs[0].runs[0].font.size = Pt(11)
        row[1].paragraphs[0].runs[0].font.bold = True
        row[2].paragraphs[0].runs[0].font.size = Pt(10)

    # Column widths
    t.columns[0].width = Inches(1.5)
    t.columns[1].width = Inches(1.3)
    t.columns[2].width = Inches(4.0)

    doc.add_paragraph()
    add_paragraph(
        doc,
        f"รวม {len(CURRENT_CROP_STAGES)} พืช — ครอบคลุม 7 category "
        "(ไม้ผล / ปาล์ม / มะพร้าว / พืชไร่ / ไม้ยืนต้น / ผักใบ / ผักผล / ถั่ว / ผักหัว)",
        italic=True, color=SUBTLE, size=10,
    )

    # --- Table 2: Standards comparison ---
    doc.add_page_break()
    add_banner(
        doc,
        "2. เทียบกับมาตรฐานไทย",
        "ตารางเทียบ vs แหล่งอ้างอิง — ต้องให้ PD ยืนยัน",
    )
    for ref in STANDARD_REFERENCES:
        add_paragraph(doc, f"▪ {ref['crop']}", bold=True, size=12, color=PRIMARY)
        add_paragraph(doc, f"   แหล่ง: {ref['source']}", size=10, color=SUBTLE, italic=True)
        add_paragraph(doc, f"   มาตรฐาน: {ref['standard']}", size=10)
        add_paragraph(doc, f"   ระยะ: {ref['stages']}", size=10)
        match_color = WARN if "ไม่" in ref['match'] or "ขาด" in ref['match'] or "ง่ายเกิน" in ref['match'] else None
        add_paragraph(doc, f"   เทียบกับของเรา: {ref['match']}",
                     size=10, bold=True, color=match_color or PRIMARY)
        doc.add_paragraph()

    # --- Table 3: Days-to-stage mapping ---
    doc.add_page_break()
    add_banner(
        doc,
        "3. Days → Stage Mapping (ยังไม่ implement)",
        "สำหรับกรณี user ตอบเป็นช่วงวัน เช่น 'หลังปลูก 30 วัน'",
    )
    add_paragraph(
        doc,
        "ตอนนี้ bot เข้าใจเฉพาะ stage name (เช่น 'ต้นกล้า', 'ออกดอก') "
        "แต่เกษตรกรจำนวนมากตอบเป็น 'อายุพืช X วัน' ซึ่งยังไม่ได้ map กลับเป็นระยะ. "
        "ข้อเสนอ mapping แบบนี้:",
        size=11,
    )
    doc.add_paragraph()

    for entry in DAYS_TO_STAGE:
        add_paragraph(doc, f"▪ {entry['crop']}", bold=True, size=12, color=PRIMARY)
        mt = doc.add_table(rows=len(entry['mapping']) + 1, cols=2)
        mt.style = "Light List Accent 1"
        c0 = mt.rows[0].cells[0]; c0.text = "อายุ"
        c1 = mt.rows[0].cells[1]; c1.text = "ระยะ"
        for c in (c0, c1):
            set_cell_shading(c, "1F6FEB")
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, (days, stage) in enumerate(entry['mapping'], 1):
            mt.rows[i].cells[0].text = days
            mt.rows[i].cells[1].text = stage
            mt.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            mt.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
            mt.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
        mt.columns[0].width = Inches(2.0)
        mt.columns[1].width = Inches(3.0)
        doc.add_paragraph()

    # --- Action items ---
    doc.add_page_break()
    add_banner(
        doc,
        "4. Action Items — สิ่งที่ต้องทำต่อ",
        "Checklist สำหรับทีม PD + dev",
    )
    action_items = [
        ("🔴", "PD review ระยะพืชทุกตัวใน section 1",
         "ยืนยันว่าคำที่ใช้ตรงกับที่เกษตรกร/dealer รู้จัก"),
        ("🔴", "PD ระบุคำเฉพาะที่ยังไม่มี",
         "เช่น ราดสาร (มะม่วง), สะสมน้ำตาล (อ้อย), ลูกขนุน (ทุเรียน)"),
        ("🟡", "ย้าย _CROP_STAGES จาก code ไป DB table crop_stages",
         "ให้ PD แก้ไขได้เองผ่าน admin panel ไม่ต้อง redeploy"),
        ("🟡", "Implement Days → Stage parser",
         "รับคำตอบแบบ 'หลังปลูก X วัน' / 'อายุ X เดือน' → map เป็น stage"),
        ("🟢", "Unit tests + E2E tests สำหรับแต่ละพืช",
         "ครอบคลุมทุกระยะ — เจอ regression ได้ทันที"),
        ("🟢", "Regional variations (ภาคเหนือ/กลาง/ใต้)",
         "ระยะบางพืช เช่น ข้าว มีความต่างตามพื้นที่"),
        ("🟢", "Symptom-to-stage inference",
         'เช่น "ใบเริ่มแดง ออกลูกแล้ว" → infer ระยะ'),
    ]
    for icon, title, desc in action_items:
        p = doc.add_paragraph()
        r = p.add_run(f"{icon} ")
        r.font.size = Pt(11)
        r = p.add_run(title)
        r.font.bold = True
        r.font.size = Pt(11)
        r = p.add_run(f"\n    {desc}")
        r.font.size = Pt(10)
        r.font.color.rgb = SUBTLE
        r.font.italic = True

    doc.add_paragraph()
    add_paragraph(
        doc,
        "🔴 = ต้องทำก่อน release  |  🟡 = ทำได้เมื่อพร้อม  |  🟢 = future enhancement",
        italic=True, color=SUBTLE, size=10,
    )

    # --- Files to update ---
    doc.add_paragraph()
    add_banner(doc, "5. ไฟล์ที่ต้องแก้ไขเมื่อ PD review แล้ว", "Reference")

    files_table = doc.add_table(rows=5, cols=3)
    files_table.style = "Light Grid Accent 1"
    rows_data = [
        ("File", "บรรทัด", "ต้องแก้ไขอะไร"),
        ("app/services/rag/response_generator_agent.py",
         "~1002-1047",
         "_CROP_STAGES dict + _DEFAULT_STAGES"),
        ("app/services/plant/registry.py",
         "~30-60",
         "เพิ่ม stage typos / synonyms ใน _TYPO_FIXES + _SYNONYMS"),
        ("migrations/create_crop_stages_table.sql",
         "(ยังไม่มี)",
         "SQL migration ถ้าย้าย _CROP_STAGES ไป DB"),
        ("tests/test_best_pick_clarify.py",
         "~9+",
         "เพิ่ม test cases ตามระยะใหม่"),
    ]
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            c = files_table.rows[i].cells[j]
            c.text = val
            if i == 0:
                set_cell_shading(c, "1F6FEB")
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                c.paragraphs[0].runs[0].font.size = Pt(10)

    # --- Footer ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(
        doc,
        "เอกสารนี้สร้างอัตโนมัติจาก scripts/generate_crop_stages_doc.py — "
        "อัพเดทล่าสุด " + datetime.now().strftime("%Y-%m-%d"),
        italic=True, size=9, color=SUBTLE,
    )

    doc.save(out_path)


def main():
    out = Path(__file__).resolve().parent.parent / "reports" / "_CROP_STAGES_DB.docx"
    out.parent.mkdir(exist_ok=True)
    build(out)
    print(f"📄 Created: {out}")


if __name__ == "__main__":
    main()
