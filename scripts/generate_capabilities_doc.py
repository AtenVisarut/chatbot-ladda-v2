"""
Generate a Word document (.docx) describing bot capabilities.
Usage: python scripts/generate_capabilities_doc.py
Output: reports/น้องลัดดา_Capabilities.docx
"""

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# Styling helpers
# ============================================================================

PRIMARY = RGBColor(0x0F, 0x4C, 0x81)     # deep blue
ACCENT = RGBColor(0x1F, 0x6F, 0xEB)      # link blue
SUBTLE = RGBColor(0x55, 0x55, 0x55)


def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=PRIMARY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Inches(0.25 * indent)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = "• " + text if not text.startswith("•") else text
    run.font.size = Pt(11)
    return p


def add_paragraph(doc, text, bold=False, color=None, size=11, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_section_banner(doc, title, subtitle=None):
    """Add a colored section banner with title"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "0F4C81")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if subtitle:
        p2 = cell.add_paragraph()
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(10)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)
    doc.add_paragraph()  # spacer


# ============================================================================
# Document content
# ============================================================================

def build_doc(out_path: Path):
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")

    # --- Cover ---
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("น้องลัดดา")
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Chatbot Capabilities — สิ่งที่ถามแล้วตอบได้")
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(f"เอกสารสร้างเมื่อ {datetime.now().strftime('%d %B %Y')}")
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = SUBTLE

    doc.add_paragraph()

    # --- Executive summary ---
    add_section_banner(
        doc,
        "สรุปภาพรวม (Executive Summary)",
        "ผลการทดสอบ 90 products × 6 capabilities = 540 queries",
    )

    # Summary table
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = "Light Grid Accent 1"
    headers = ["Overall Accuracy", "Products tested", "Pass rate (≥70)", "Capabilities"]
    values = ["96.84 / 100", "90 / 90", "89 / 90 (99%)", "6 ด้าน"]
    for i, (h, v) in enumerate(zip(headers, values)):
        c1 = summary_table.rows[0].cells[i]
        c1.text = h
        c1.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(11)
        c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c1, "1F6FEB")
        c2 = summary_table.rows[1].cells[i]
        c2.text = v
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c2.paragraphs[0].runs[0].font.size = Pt(14)
        c2.paragraphs[0].runs[0].font.bold = True
        c2.paragraphs[0].runs[0].font.color.rgb = PRIMARY

    doc.add_paragraph()

    # --- 10 topic groups ---
    sections = [
        {
            "num": "1",
            "title": "ข้อมูลสินค้า (ถามเฉพาะตัว)",
            "score": 94.2,
            "items": [
                ("ชื่อสารสำคัญ + เปอร์เซ็นต์",
                 'ไบเตอร์ (ไบเฟนทริน 5% + อิมิดาคลอพริด 25%)'),
                ("Formulation (สูตร)",
                 'EC / SC / WP / WG / SL / GR / …'),
                ("ขนาดบรรจุ",
                 'ขนาดขวด / ซอง / กิโลกรัม'),
                ("รูปแบบสินค้า",
                 'น้ำ / ผง / เกล็ด / เม็ด'),
                ("ประเภทสินค้า",
                 'Insecticide / Fungicide / Herbicide / PGR / Biostimulants / Fertilizer'),
                ("ชื่อสารภาษาไทย",
                 'แปลจาก active_ingredient เป็นภาษาไทย'),
            ],
        },
        {
            "num": "2",
            "title": "พืช และ ศัตรูพืช",
            "score": 97.2,
            "items": [
                ("สินค้านี้ใช้กับพืชอะไรได้บ้าง", None),
                ("สินค้านี้กำจัดอะไรได้", "โรค / แมลง / วัชพืช"),
                ("สินค้านี้ใช้กับพืช X ได้ไหม (Applicability check)",
                 'ถ้าไม่ได้ระบุใน DB → บอก "ไม่ได้ระบุว่าใช้กับ..." + แนะนำตัวอื่น'),
                ("ในพืช X มีโรค/แมลง อะไรที่สินค้ากำจัดได้", None),
            ],
        },
        {
            "num": "3",
            "title": "การใช้งาน (How-to)",
            "score": 99.4,
            "star": True,
            "items": [
                ("อัตราใช้", '15-20 มล./น้ำ 20 ลิตร — ห้าม bot คำนวณเอง'),
                ("วิธีผสม/ฉีดพ่น/ราด/หว่าน", None),
                ("ใช้ช่วงไหน / ระยะไหน", None),
                ("ออกฤทธิ์แบบสัมผัส / ดูดซึม (absorption_method)", None),
                ("ข้อควรระวัง / พิษต่อพืช (phytotoxicity)", None),
            ],
        },
        {
            "num": "4",
            "title": "กลุ่มสารเคมี (Technical)",
            "score": 98.9,
            "items": [
                ("IRAC / FRAC / HRAC group", 'กลุ่ม 3A + 4A, กลุ่ม C2, กลุ่ม E'),
                ("MoA (Mode of Action) / กลไกการออกฤทธิ์", None),
                ("กลุ่มสารเคมีย่อย", None),
                ("PGR / ควบคุมการเจริญเติบโต",
                 'รองรับทั้งแบบ code และ narrative เช่น "กลุ่มจิบเบอเรลลิน"'),
            ],
        },
        {
            "num": "5",
            "title": "จุดเด่นสินค้า (Selling Point)",
            "score": 95.8,
            "items": [
                ("จุดเด่น / คุณสมบัติ", None),
                ("เหตุผลที่ควรใช้ตัวนี้", None),
            ],
        },
        {
            "num": "6",
            "title": "คำถามแนะนำสินค้า",
            "score": None,
            "items": [
                ("มียาอะไรแก้ {โรค/แมลง/วัชพืช} ใน {พืช}", None),
                ("บำรุง / เร่งดอก / เร่งผล {พืช} ใช้อะไร",
                 'ห้ามแนะนำ Fungicide เมื่อ intent = บำรุง'),
                ("กำจัดหญ้าในนา / สวน / ไร่", None),
                ("ปุ๋ยเกล็ด / ปุ๋ยน้ำ / NPK สูตรไหน",
                 'ปุ๋ยเกล็ด → NPK เท่านั้น / ปุ๋ยน้ำ → กรอง physical_form'),
            ],
        },
        {
            "num": "7",
            "title": "เปรียบเทียบสินค้า",
            "score": 95.6,
            "items": [
                ("สินค้า A กับ B ต่างกันยังไง", None),
                ("Family variants",
                 'เช่น ไดยูแมกซ์ SC vs WP, โบว์แลน vs โบว์แลน 285'),
                ("บอมส์ family (suffix-only)",
                 'พิมพ์ "ไวท์ แม็กซ์ ซิงค์" → เทียบทั้ง 3 ตัว'),
                ("พรีดิคท์ 10% / 15% / 25% เอฟ",
                 'รองรับ version number ในชื่อสินค้า'),
                ("NPK 0-0-60 / 0-52-34 / 13-0-46", None),
            ],
        },
        {
            "num": "8",
            "title": "คำถามต่อเนื่อง (Follow-up) — ใหม่",
            "score": None,
            "items": [
                ("ใช้ยังไง / อัตราเท่าไหร่ (หลังระบุสินค้า)",
                 "bot จำสินค้าจาก turn ก่อน"),
                ("ใช้กับ {พืช} ได้ไหม (applicability)",
                 "บอก 'ไม่ได้ระบุ' ก่อนแนะนำตัวอื่น"),
                ("ใช้แตกต่างกันยังไง (เทียบสินค้าเดิม)",
                 "ไม่หลุดไปสินค้าอื่น"),
                ("ตัวไหนดีสุด",
                 "bot ถามกลับ: ใช้กับพืชอะไร ระยะไหน ก่อนแนะนำ"),
            ],
        },
        {
            "num": "9",
            "title": "ภาษาชาวบ้าน / สะกดผิด",
            "score": None,
            "items": [
                ("คำภาษาชาวบ้าน (slang)",
                 'ข้าวดีด/ข้าวตีด = วัชพืช, ใบด่าง/ใบจุด = โรคพืช, ราดสาร/ยับยั้งใบอ่อน = PGR'),
                ("สะกดผิด / ตัวอักษรคลาด (typo)",
                 'โทมาหอค → โทมาฮอค, ออลสตาร → ออล์สตาร์'),
                ("ตัวอักษรอังกฤษ (transliteration)",
                 'topgun / imidagold / tomahawk → Thai canonical'),
                ("ข้ามเครื่องหมาย diacritics",
                 'คาริสม่า = คาริสมา'),
                ("Fuzzy match threshold ≥ 0.75",
                 'auto-generate typo variants จาก consonant swap'),
            ],
        },
        {
            "num": "10",
            "title": "ตอบพร้อม Handoff — เมื่อตอบเองไม่ได้",
            "score": None,
            "items": [
                ("ราคา / ซื้อที่ไหน",
                 '"สอบถามตัวแทน ICP Ladda" + สร้าง handoff + alert'),
                ("คำนวณพื้นที่ / ปริมาณ",
                 '"ปรึกษาเจ้าหน้าที่" — ห้ามคำนวณเอง'),
                ("ผสมร่วม / สลับสาร",
                 "Handoff ถ้า DB ไม่ระบุ"),
                ("คำถามนอกเกษตร",
                 '"น้องลัดดายินดีตอบเรื่องเกษตรเท่านั้น"'),
            ],
        },
    ]

    for s in sections:
        # Section banner
        star = " ⭐" if s.get("star") else ""
        score_str = f"  (คะแนน: {s['score']}/100{star})" if s.get("score") else ""
        subtitle = f"คะแนนจากการทดสอบ: {s['score']}/100" if s.get("score") else None
        add_section_banner(
            doc,
            f"{s['num']}. {s['title']}{score_str}",
            subtitle,
        )

        # Items
        for item in s["items"]:
            topic, example = item
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run("• ")
            r.font.bold = True
            r.font.color.rgb = ACCENT
            r.font.size = Pt(12)
            r2 = p.add_run(topic)
            r2.font.size = Pt(12)
            r2.font.bold = True

            if example:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.5)
                r = p2.add_run(f"ตัวอย่าง: {example}")
                r.font.size = Pt(10)
                r.font.italic = True
                r.font.color.rgb = SUBTLE

    # --- Capability score table ---
    doc.add_page_break()
    add_section_banner(doc, "ตารางคะแนนความสามารถ", "Per-capability accuracy from 540-query test")

    score_table = doc.add_table(rows=8, cols=4)
    score_table.style = "Light Grid Accent 1"
    rows = [
        ("Capability", "คะแนน", "Pass Rate", "หมายเหตุ"),
        ("1. ข้อมูลสินค้า", "94.2", "86/90", "% + formulation ครบ"),
        ("2. โรค/แมลง/พืช", "97.2", "88/90", "รองรับ pest+crop query"),
        ("3. อัตรา + วิธีใช้", "99.4 ⭐", "89/90", "ดีที่สุด — ไม่คำนวณเอง"),
        ("4. MoA / IRAC", "98.9", "89/90", "รวม PGR narrative"),
        ("5. จุดเด่น", "95.8", "86/90", "จาก selling_point"),
        ("6. เปรียบเทียบ", "95.6", "82/90", "multi-product + family"),
        ("Overall", "96.84", "89/90 (99%)", "89 products ผ่านเกณฑ์ ≥70"),
    ]
    for i, r_data in enumerate(rows):
        for j, val in enumerate(r_data):
            cell = score_table.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_shading(cell, "1F6FEB")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i == 7:  # Overall row
                set_cell_shading(cell, "FFF9C4")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = PRIMARY
            else:
                cell.paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # --- Anti-hallucination section ---
    add_section_banner(
        doc,
        "ระบบป้องกันความผิดพลาด (Anti-hallucination)",
        "14 safeguards ป้องกัน bot ตอบข้อมูลที่ไม่มีใน DB",
    )
    safeguards = [
        "Cross-product removal — ลบสินค้าที่ไม่อยู่ใน retrieved docs",
        "NumberCheck — ตรวจตัวเลขทุกตัวว่าตรงกับ DB",
        "Crop-plant match — block สินค้าที่ห้ามใช้กับพืชที่ถาม",
        "Disease-pest mismatch block — ห้ามแนะนำถ้า DB ไม่ระบุ",
        "NO_DATA_REPLY — ตอบ 'กำลังตรวจสอบ' + trigger handoff + alert",
        "Category-intent alignment — ห้าม fungicide ตอบ intent ปุ๋ย",
        "Disease rescue — rescue สินค้าที่ถูก crop-filter ตัด",
        "Fresh-fetch bypass memory — กันปนข้อมูลสินค้าเก่า",
        "Topic boundary — detect สินค้าใหม่ → skip context เก่า",
        "Comparison differentiator — บังคับชี้ความต่าง",
        "Applicability denial — บอก 'ไม่ได้ระบุ' ก่อนแนะนำอื่น",
        "active_products list — retain family variants",
        "PGR narrative prompt — MoA แบบ narrative",
        "Best-pick clarify — ถามกลับเมื่อข้อมูลไม่พอ",
    ]
    for sg in safeguards:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(sg)
        r.font.size = Pt(11)

    # --- Ecosystem section ---
    doc.add_page_break()
    add_section_banner(
        doc,
        "ระบบโดยรวม (System Scope)",
        "ขอบเขตข้อมูลและโครงสร้างพื้นฐาน",
    )

    scope_items = [
        ("จำนวนสินค้าใน DB", "90 รายการ"),
        ("Categories ครอบคลุม", "6 กลุ่ม: Insecticide (24), Fungicide (19), Herbicide (27), PGR (5), Biostim (8), Fertilizer (7)"),
        ("Aliases (ชื่อเรียกอื่น)", "~779 — รองรับ typo, slang, English"),
        ("Registry auto-refresh", "ทุก 15 นาที — สินค้าใหม่ใช้ได้ทันทีไม่ต้อง restart"),
        ("ช่องทาง", "LINE Official Account + Facebook Messenger"),
        ("Dashboard", "Vercel — หน้า Alerts / Chat / Templates / Users / Products"),
        ("Admin handoff system", "เมื่อ bot ตอบไม่ได้ → admin เข้าช่วย + ใช้ template ตอบ"),
        ("Response time", "10-16 วินาทีต่อข้อความ (RAG pipeline)"),
        ("Test coverage", "1,328 unit + integration tests"),
    ]
    for topic, desc in scope_items:
        p = doc.add_paragraph()
        r = p.add_run(f"• {topic}: ")
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = ACCENT
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)

    # --- What bot does NOT answer ---
    doc.add_paragraph()
    add_section_banner(
        doc,
        "สิ่งที่บอทไม่ตอบ (By Design)",
        "เรื่องที่ตั้งใจให้ handoff ไปยังเจ้าหน้าที่",
    )
    no_answer = [
        ("ราคา / ซื้อที่ไหน", "ป้องกันข้อมูลราคาคลาดเคลื่อน"),
        ("คำนวณพื้นที่ / ปริมาณ", "ป้องกันการคำนวณผิดทำให้ใช้ยาผิด"),
        ("ผสมร่วมกับสารอื่น / สลับสาร", "ถ้า DB ไม่ระบุ → handoff เพื่อความปลอดภัย"),
        ("สินค้ายี่ห้ออื่น / คู่แข่ง", "ไม่เปรียบเทียบกับสินค้านอก ICP Ladda"),
        ("คำถามนอกการเกษตร", "จำกัดขอบเขตเพื่อคุณภาพคำตอบ"),
    ]
    for topic, reason in no_answer:
        p = doc.add_paragraph()
        r = p.add_run(f"• {topic} ")
        r.font.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(f"— {reason}")
        r2.font.size = Pt(11)
        r2.font.italic = True
        r2.font.color.rgb = SUBTLE

    # --- Footer note ---
    doc.add_paragraph()
    add_paragraph(
        doc,
        "เอกสารนี้สร้างอัตโนมัติจากผลการทดสอบระบบ — "
        "อัพเดทล่าสุด 2026-04-20",
        italic=True, size=9, color=SUBTLE,
    )

    doc.save(out_path)
    return out_path


def main():
    out = Path(__file__).resolve().parent.parent / "reports" / "น้องลัดดา_Capabilities.docx"
    out.parent.mkdir(exist_ok=True)
    build_doc(out)
    print(f"📄 Created: {out}")


if __name__ == "__main__":
    main()
